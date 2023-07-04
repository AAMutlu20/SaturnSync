import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def get_website_data(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    
    title = soup.title.string if soup.title else ''
    
    paragraphs = []
    for paragraph in soup.find_all('p'):
        text = paragraph.get_text().strip()
        if text:
            paragraphs.append(text)
    
    image_urls = []
    for img in soup.find_all('img'):
        img_url = img.get('src')
        if img_url:
            image_urls.append(urljoin(url, img_url))
    
    return title, sorted(paragraphs), image_urls

def save_to_txt(title, paragraphs, image_urls, file_path):
    with open(file_path, 'w') as file:
        file.write(f'Title: {title}\n\n')
        
        file.write('Text Data:\n\n')
        for paragraph in paragraphs:
            file.write(paragraph + '\n\n')
        
        file.write('Image URLs:\n\n')
        for img_url in image_urls:
            file.write(img_url + '\n')

def save_to_word(title, paragraphs, image_urls, file_path):
    doc = Document()

    title_heading = doc.add_heading(level=1)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_heading.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Arial'
    
    doc.add_paragraph()
    
    text_heading = doc.add_heading('Text Data', level=2)
    text_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    text_heading.runs[0].bold = True
    text_heading.runs[0].font.size = Pt(12)
    text_heading.runs[0].font.name = 'Arial'
    
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph, style='BodyText')
    
    doc.add_paragraph()

    image_heading = doc.add_heading('Image URLs', level=2)
    image_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    image_heading.runs[0].bold = True
    image_heading.runs[0].font.size = Pt(12)
    image_heading.runs[0].font.name = 'Arial'

    for img_url in image_urls:
        doc.add_paragraph(img_url, style='BodyText')

    doc.save(file_path)

def display_website_data(url):
    title, paragraphs, image_urls = get_website_data(url)
    
    print('Title:')
    print(title)
    print()
    
    print('Text Data:')
    for paragraph in paragraphs:
        print(paragraph)
        print()
    
    print('Image URLs:')
    for img_url in image_urls:
        print(img_url)

    save_option = input("Save the data to a file? Enter 'txt', 'word', or 'none': ")
    if save_option == 'txt':
        file_path = input("Enter the file path to save the TXT file: ")
        save_to_txt(title, paragraphs, image_urls, file_path)
        print("Data saved to TXT file.")
    elif save_option == 'word':
        file_path = input("Enter the file path to save the Word file: ")
        save_to_word(title, paragraphs, image_urls, file_path)
        print("Data saved to Word file.")
    elif save_option == 'none':
        print("Data not saved.")
    else:
        print("Invalid option. Data not saved.")

website_url = input("Enter the website URL: ")
display_website_data(website_url)